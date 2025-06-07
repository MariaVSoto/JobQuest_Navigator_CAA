import PyPDF2
import docx
import nltk
import requests
import spacy
import openai
import os
from bs4 import BeautifulSoup
from nltk.tokenize import wordpunct_tokenize
from nltk.corpus import stopwords

# Download necessary resources
nltk.download('punkt')
nltk.download('stopwords')

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

print(openai.__version__)

class ResumeAnalyzer:
    def __init__(self):
        self.stop_words = set(stopwords.words('english'))
        self.nlp = nlp

    def extract_text_from_pdf(self, file_path):
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                total_pages = len(reader.pages)
                print(f"Processing PDF with {total_pages} pages...")
                
                for i, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text()
                    if not page_text.strip():
                        print(f"Warning: Page {i} appears to be empty or could not be extracted properly")
                    # Preserve bullet points and formatting
                    page_text = page_text.replace('•', '\n•').replace('·', '\n·')
                    # Add page number as separator
                    text += f"\n=== PAGE {i} ===\n{page_text}\n"
                    print(f"Processed page {i}/{total_pages}")
                
                if not text.strip():
                    raise ValueError("No text could be extracted from the PDF")
                
                print(f"Successfully extracted {len(text.split())} words from the PDF")
                return text
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            raise

    def extract_text_from_doc(self, file_path):
        try:
            doc = docx.Document(file_path)
            text = ""
            total_paragraphs = len(doc.paragraphs)
            print(f"Processing DOCX with {total_paragraphs} paragraphs...")
            
            for i, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                if i % 10 == 0:  # Print progress every 10 paragraphs
                    print(f"Processed {i}/{total_paragraphs} paragraphs")
            
            if not text.strip():
                raise ValueError("No text could be extracted from the DOCX file")
            
            print(f"Successfully extracted {len(text.split())} words from the DOCX")
            return text
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            raise

    def extract_job_description(self, url):
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        # Try to find a main job description div, fallback to all text
        job_description_div = soup.find('div', class_='job-description')
        if job_description_div:
            return job_description_div.text
        else:
            return soup.get_text()

    def analyze_resume(self, resume_text, job_description):
        # Process text with spaCy
        resume_doc = self.nlp(resume_text)
        job_doc = self.nlp(job_description)

        # Define all possible sections with their variations
        section_keywords = {
            'summary': ['summary', 'profile', 'professional summary', 'career summary', 'objective', 'career objective'],
            'skills': ['core skills', 'skills', 'technical skills', 'key skills', 'competencies', 'expertise', 'capabilities'],
            'experience': ['experience', 'work experience', 'employment', 'professional experience', 'work history', 'employment history'],
            'projects': ['projects', 'project experience', 'portfolio', 'project portfolio', 'selected projects', 'key projects'],
            'education': ['education', 'academic', 'qualifications', 'academic background', 'educational background'],
            'certifications': ['certifications', 'certificates', 'professional certifications', 'licenses'],
            'languages': ['languages', 'language skills', 'language proficiency'],
            'other': []  # Default section for unmatched content
        }
        
        # Initialize sections dictionary
        sections = {section: [] for section in section_keywords.keys()}
        current_section = 'other'
        
        # Split text into lines and process
        lines = resume_text.split('\n')
        buffer = []  # Buffer to store lines until we determine their section
        
        for line in lines:
            line = line.strip()
            if not line:  # Skip empty lines
                continue
            
            line_lower = line.lower()
            
            # Check if this line indicates a new section
            found_section = False
            for section, keywords in section_keywords.items():
                if any(keyword in line_lower for keyword in keywords):
                    # If we have buffered content, add it to the previous section
                    if buffer:
                        sections[current_section].extend(buffer)
                        buffer = []
                    current_section = section
                    found_section = True
                    break
            
            # If this line didn't indicate a new section, add it to buffer
            if not found_section:
                buffer.append(line)
        
        # Add any remaining buffered content to the current section
        if buffer:
            sections[current_section].extend(buffer)
        
        # Process each section
        resume_entities = []
        resume_tokens = []
        
        for section_name, section_text in sections.items():
            if section_text:  # Only process non-empty sections
                section_doc = self.nlp('\n'.join(section_text))
                section_entities = [ent.text for ent in section_doc.ents]
                section_tokens = wordpunct_tokenize('\n'.join(section_text).lower())
                section_tokens = [word for word in section_tokens if word.isalnum() and word not in self.stop_words]
                
                resume_entities.extend(section_entities)
                resume_tokens.extend(section_tokens)
                
                print(f"\nProcessed {section_name} section:")
                print(f"Found {len(section_entities)} entities and {len(section_tokens)} keywords")
                print(f"Section content preview: {' '.join(section_text[:3])}...")

        # Extract job requirements
        job_entities = [ent.text for ent in job_doc.ents]
        job_tokens = wordpunct_tokenize(job_description.lower())
        job_tokens = [word for word in job_tokens if word.isalnum() and word not in self.stop_words]

        # Find missing elements
        missing_entities = set(job_entities) - set(resume_entities)
        missing_keywords = set(job_tokens) - set(resume_tokens)

        return list(missing_entities), list(missing_keywords), sections

    def generate_suggestions(self, resume_text, job_description):
        missing_entities, missing_keywords, sections = self.analyze_resume(resume_text, job_description)
        
        suggestions = []
        
        # Experience section suggestions
        if sections['experience']:
            exp_text = '\n'.join(sections['experience'])
            exp_doc = self.nlp(exp_text)
            exp_entities = [ent.text for ent in exp_doc.ents]
            exp_tokens = wordpunct_tokenize(exp_text.lower())
            exp_tokens = [word for word in exp_tokens if word.isalnum() and word not in self.stop_words]
            
            missing_exp_entities = set(job_entities) - set(exp_entities)
            missing_exp_keywords = set(job_tokens) - set(exp_tokens)
            
            if missing_exp_entities or missing_exp_keywords:
                suggestions.append("Experience Section:")
                if missing_exp_entities:
                    suggestions.append(f"- Consider adding these skills/qualifications to your experience: {', '.join(missing_exp_entities)}")
                if missing_exp_keywords:
                    suggestions.append(f"- Include these keywords in your experience descriptions: {', '.join(missing_exp_keywords)}")
        
        # Projects section suggestions
        if sections['projects']:
            proj_text = '\n'.join(sections['projects'])
            proj_doc = self.nlp(proj_text)
            proj_entities = [ent.text for ent in proj_doc.ents]
            proj_tokens = wordpunct_tokenize(proj_text.lower())
            proj_tokens = [word for word in proj_tokens if word.isalnum() and word not in self.stop_words]
            
            missing_proj_entities = set(job_entities) - set(proj_entities)
            missing_proj_keywords = set(job_tokens) - set(proj_tokens)
            
            if missing_proj_entities or missing_proj_keywords:
                suggestions.append("\nProjects Section:")
                if missing_proj_entities:
                    suggestions.append(f"- Consider adding these technologies/tools to your projects: {', '.join(missing_proj_entities)}")
                if missing_proj_keywords:
                    suggestions.append(f"- Include these keywords in your project descriptions: {', '.join(missing_proj_keywords)}")
        
        # General suggestions
        suggestions.append("\nGeneral Recommendations:")
        suggestions.append("- Ensure your experience and projects are presented in reverse chronological order")
        suggestions.append("- Quantify your achievements with specific numbers and metrics")
        suggestions.append("- Use action verbs to start each bullet point")
        suggestions.append("- Keep descriptions concise and impactful")
        
        return suggestions

class UIReview:
    def __init__(self):
        pass

    def display_suggestions(self, suggestions):
        for suggestion in suggestions:
            print(f"Suggestion: {suggestion}")

    def accept_suggestion(self, suggestion):
        print(f"Accepted: {suggestion}")

    def reject_suggestion(self, suggestion):
        print(f"Rejected: {suggestion}")

class FeedbackLoop:
    def __init__(self):
        self.feedback_data = []

    def log_feedback(self, suggestion, action):
        self.feedback_data.append((suggestion, action))
        print(f"Logged: {suggestion} - {action}")

    def update_model(self):
        print("Model updated based on feedback.")

def get_ai_resume_review(resume_text, job_description, openai_api_key):
    client = openai.OpenAI(api_key=openai_api_key)
    prompt = f"""
You are an expert career coach and resume reviewer with deep knowledge of ATS systems and hiring practices.
Here is a resume:
---
{resume_text}
---
And here is a job description:
---
{job_description}
---
Please provide a detailed, line-by-line review of the resume with the following:

1. Review each section in the order it appears in the resume:
   - Quote the COMPLETE original section
   - Provide specific, actionable suggestions for improvement
   - Include concrete examples
   - Quantify achievements where possible
   - Ensure suggestions align with the job requirements

2. For each section, follow this format:
   SECTION: [Name of the section as it appears in the resume]
   ORIGINAL CONTENT: [Quote the complete section]
   SUGGESTED CONTENT: [Show the complete suggested text]
   REASONING: [Explain why this improves the resume]
   IMPACT: [How this helps match the job requirements]

3. For each suggestion:
   - Show the COMPLETE original text
   - Show the COMPLETE suggested text
   - Explain WHY the change would improve the resume
   - Include specific metrics or numbers when relevant
   - Highlight how it aligns with the job requirements

4. Important guidelines:
   - Review ALL content from ALL pages
   - Process sections in the order they appear in the resume
   - Provide COMPLETE suggestions for each section
   - Include specific examples
   - Show before and after text
   - Explain the impact on job requirements
   - Never leave phrases or suggestions incomplete

5. End with an overall assessment and priority of changes.

IMPORTANT: 
- Always show complete sentences and paragraphs
- Review ALL content from ALL pages
- Process sections in their original order
- Provide specific, actionable improvements
- Focus on quality over quantity of suggestions
- Ensure suggestions are relevant to the job requirements
- Pay special attention to the first section (usually Summary/Profile)
- Make sure to review the last section completely

Write your suggestions in clear, professional English, focusing on actionable improvements.
"""
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4000,
        temperature=0.7
    )
    return response.choices[0].message.content.strip()

def main():
    # Get OpenAI API key
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        openai_api_key = input("Enter your OpenAI API key: ").strip()

    analyzer = ResumeAnalyzer()
    ui = UIReview()
    feedback = FeedbackLoop()

    # Ask user for resume file
    file_path = input("Enter the path to your resume (PDF or DOCX): ")
    if file_path.endswith('.pdf'):
        resume_text = analyzer.extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        resume_text = analyzer.extract_text_from_doc(file_path)
    else:
        print("Unsupported file format. Please use PDF or DOCX.")
        return

    # Ask user for job URL
    job_url = input("Enter the URL of the job position: ")
    job_description = analyzer.extract_job_description(job_url)

    # Generate AI review
    print("\nGenerating AI-powered resume review...\n")
    review = get_ai_resume_review(resume_text, job_description, openai_api_key)
    print("AI Resume Review:\n")
    print(review)

    # Save the review to a file with the same name as input but with _review suffix
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_path = f"{base_name}_review.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(review)
    print(f"\nReview saved to {output_path}")

if __name__ == "__main__":
    main()